import os
import yaml
import re
import glob
import hashlib
from pathlib import Path
from playwright.sync_api import sync_playwright

_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_W_P = _W_NS + "p"
_W_R = _W_NS + "r"
_W_T = _W_NS + "t"
_W_UPDATE_FIELDS = _W_NS + "updateFields"
_W_VAL = _W_NS + "val"
_V_NS = "{urn:schemas-microsoft-com:vml}"
_R_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def _strip_orphan_vml_shapes(doc):
    """删除内容文档里没有有效 r:id 的 <v:shape>。
    docxcompose.insert -> add_shapes 在迁移图片关系时会查 doc.part.rels[rid]，
    rid 为空就抛 KeyError(None)。这类孤立 VML 形状不是正文真正的图片
    （真正的图片走 <w:drawing>），直接清理即可。"""
    removed = 0
    for shape in list(doc.element.iter(_V_NS + "shape")):
        bad = False
        imagedatas = list(shape.iter(_V_NS + "imagedata"))
        if not imagedatas:
            continue
        for im in imagedatas:
            rid = im.get(_R_NS + "id")
            if not rid:
                bad = True
                break
        if bad:
            parent = shape.getparent()
            if parent is not None:
                parent.remove(shape)
                removed += 1
    return removed


def _force_toc_refresh_on_open(doc):
    """往 settings.xml 注入 <w:updateFields w:val="true"/>，
    让 Word 打开文档时自动刷新 TOC/页码等字段。"""
    from lxml import etree
    settings_el = doc.settings.element
    existing = settings_el.find(_W_UPDATE_FIELDS)
    if existing is not None:
        existing.set(_W_VAL, "true")
        return
    el = etree.SubElement(settings_el, _W_UPDATE_FIELDS)
    el.set(_W_VAL, "true")


def _own_t_nodes(p):
    """返回直接属于段落 p、不在 p 的嵌套子段落里的 <w:t> 节点列表。

    这一步很关键：当 p 内嵌 <w:drawing> 包文本框时，文本框里又有自己的
    <w:p>。`p.iter(_W_T)` 会把子段落里的 <w:t> 一起带出来，造成"跨段聚合"，
    导致同一段被替换多次（封面 Choice/Fallback 两份内容被合并成一段）。
    这里手动校验每个 <w:t> 的最近 w:p 祖先必须等于 p 本身。
    """
    out = []
    for t in p.iter(_W_T):
        anc = t.getparent()
        while anc is not None and anc.tag != _W_P:
            anc = anc.getparent()
        if anc is p:
            out.append(t)
    return out


def _replace_placeholder_everywhere(root_element, placeholder, value):
    """替换 root_element 下所有 <w:p> 段落里的占位符（含文本框 <w:txbxContent>、表格、形状）。
    占位符可能被拆分到多个 <w:r>/<w:t> 中，因此按段落聚合文本后写回首个 <w:t>，
    其余 <w:t> 清空，保留首个 run 的字号/字体样式。

    注意：仅聚合"本段直属"的 <w:t>，避免外层段落把内嵌文本框子段落的内容
    也吞并进来，否则封面 mc:Choice + mc:Fallback 两份内容会被拼成双份。
    """
    variants = (placeholder, placeholder.replace(" ", ""))
    replaced = 0
    for p in root_element.iter(_W_P):
        t_nodes = _own_t_nodes(p)
        if not t_nodes:
            continue
        full_text = "".join(t.text or "" for t in t_nodes)
        if not any(v in full_text for v in variants):
            continue
        new_text = full_text
        for v in variants:
            new_text = new_text.replace(v, value)
        t_nodes[0].text = new_text
        for t in t_nodes[1:]:
            t.text = ""
        replaced += 1
    return replaced


def load_config(config_path="export-map.yaml"):
    if not os.path.exists(config_path):
        print(f"错误: 找不到配置文件 {config_path}")
        return None
    with open(config_path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)

def take_screenshot(html_path, output_path, action=None):
    # 使用 pathlib 的 as_uri()，完美解决 Windows 下的 file:///C:/... 协议问题
    abs_html_path = Path(html_path).resolve().as_uri()
    print(f"正在截图: {html_path} -> {output_path}")
    if action:
         print(f"  执行动作: {action}")
    
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        # 设置一个较大的视口以确保截取全屏，并设置 device_scale_factor 为 3.0 获取极致高清
        page = browser.new_page(viewport={"width": 1920, "height": 1080}, device_scale_factor=3.0)
        page.goto(abs_html_path)
        # 等待页面加载完成
        page.wait_for_load_state('networkidle')
        
        # 如果配置了交互动作（例如打开弹窗）
        if action:
            if action.startswith('click:'):
                # 示例: "click:button:has-text('新建')" 或 "click:#newBtn"
                selector = action.split('click:', 1)[1]
                try:
                    # 尝试定位并点击，由于可能存在多个同名按钮（例如表格每一行都有编辑），默认点击第一个
                    page.locator(selector).first.click()
                    # 等待动画完成，抽屉/弹窗的过渡动画通常需要 300-500ms，这里放宽到 800ms
                    page.wait_for_timeout(800) 
                except Exception as e:
                    print(f"  执行动作失败 {action}: {e}")

        # 全屏截图
        page.screenshot(path=output_path, full_page=True)
        browser.close()

def process_markdown(md_path, html_path, screenshot_path, replace_keyword):
    # 因为一个 md 文件可能被多次处理（多次截图插入不同位置）
    # 我们应该读取最新的内容（如果存在临时文件的话）
    temp_md_path = md_path + ".tmp"
    target_read_path = temp_md_path if os.path.exists(temp_md_path) else md_path

    print(f"正在处理 Markdown: {target_read_path} (寻找关键字: {replace_keyword})")
    with open(target_read_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 使用绝对路径，并添加 width=100% 确保在 Word 中撑满页面且不失真
    image_md = f"\n![页面原型]({os.path.abspath(screenshot_path)}){{width=100%}}\n"

    # 注意：replace_keyword 本身已经包含了正则语法（如 "[:：]\s*"），不要再用 re.escape 转义它
    # 模式：找到关键字，然后匹配中间的任意内容（非贪婪），直到遇到 ```text 块，并捕获它
    # 尝试匹配 3 个变种，解决 "页面名称: xxx" 与 ```text 之间可能夹杂任何奇怪字符的问题
    # 变种1：紧跟
    pattern1 = re.compile(rf"({replace_keyword}.*?)(```text[\s\S]*?```)", re.IGNORECASE | re.DOTALL)
    # 变种2：如果 Markdown 里的 `页面名称` 没有被冒号包裹，而是比如 `### 1.1 车辆管理列表`
    # 提取真正的纯文字页面名
    clean_keyword = re.sub(r'页面名称\[:：\]\\s\*', '', replace_keyword)
    clean_keyword = clean_keyword.replace('[:：]\\s*', '')
    
    pattern2 = re.compile(rf"({clean_keyword}.*?)(```text[\s\S]*?```)", re.IGNORECASE | re.DOTALL)
    
    if pattern1.search(content):
        new_content = pattern1.sub(rf"\1{image_md}", content, count=1)
        with open(temp_md_path, 'w', encoding='utf-8') as f:
            f.write(new_content)
        return temp_md_path
    elif pattern2.search(content):
        new_content = pattern2.sub(rf"\1{image_md}", content, count=1)
        with open(temp_md_path, 'w', encoding='utf-8') as f:
            f.write(new_content)
        return temp_md_path
    else:
        # 宽松模式：只要找到纯文字名字，就替换它后面的第一个 ```text
        loose_pattern = re.compile(rf"({clean_keyword}.{{0,500}}?)(```text[\s\S]*?```)", re.IGNORECASE | re.DOTALL)
        if loose_pattern.search(content):
             new_content = loose_pattern.sub(rf"\1{image_md}", content, count=1)
             with open(temp_md_path, 'w', encoding='utf-8') as f:
                 f.write(new_content)
             return temp_md_path
             
        print(f"警告: 在文档中未找到指定的关键字 '{replace_keyword}' (纯文本: '{clean_keyword}') 下的 ASCII 线框图。")
        return target_read_path

def render_mermaid_blocks(content, screenshot_dir):
    """
    寻找 Markdown 中的 mermaid 代码块，使用 Playwright 渲染为图片并替换原文。
    """
    pattern = re.compile(r'```mermaid[ \t]*\n(.*?)```', re.DOTALL)
    matches = pattern.findall(content)
    
    if not matches:
        return content
        
    print(f"发现 {len(matches)} 个 Mermaid 图表，正在渲染为图片...")
    
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        
        def repl(match):
            mermaid_code = match.group(1)
            # 使用 hash 生成唯一文件名
            code_hash = hashlib.md5(mermaid_code.encode('utf-8')).hexdigest()[:8]
            img_name = f"mermaid_{code_hash}.png"
            img_path = os.path.join(screenshot_dir, img_name)
            
            # 对于 Mermaid 代码，有些符号在 HTML 中可能会导致解析错误
            # 更好的做法是不直接放在 div 中，而是通过 JavaScript 动态插入
            # 为了防止引号导致的问题，将代码 base64 编码后在 JS 中解码
            import base64
            encoded_mermaid = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
            
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
              <script type="module">
                import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
                mermaid.initialize({{ startOnLoad: false, theme: 'default' }});
                
                window.onload = async () => {{
                    const code = decodeURIComponent(escape(window.atob('{encoded_mermaid}')));
                    const container = document.getElementById('mermaid-container');
                    const errorContainer = document.getElementById('error-container');
                    try {{
                        const {{ svg }} = await mermaid.render('mermaid-svg', code);
                        container.innerHTML = svg;
                    }} catch (e) {{
                        console.error('Mermaid render error:', e);
                        errorContainer.innerHTML = `<div style="color:red;font-family:sans-serif;font-size:24px;border:2px solid red;padding:20px;">Syntax error in Mermaid:<br><pre>${{e.message || e}}</pre><br>Code:<br><pre>${{code}}</pre></div>`;
                        container.classList.add('error');
                    }}
                }};
              </script>
              <style>
                body {{ background: white; padding: 20px; display: inline-block; margin: 0; }}
              </style>
            </head>
            <body>
              <div id="mermaid-container" class="mermaid"></div>
              <div id="error-container"></div>
            </body>
            </html>
            """
            temp_html = os.path.abspath(f"temp_mermaid_{code_hash}.html")
            with open(temp_html, "w", encoding="utf-8") as f:
                f.write(html_content)
            
            try:
                # 使用极高的 device_scale_factor 获取印刷级 5 倍超清 PNG
                page = browser.new_page(device_scale_factor=5.0)
                # 使用 pathlib 确保跨平台 file 协议兼容
                page.goto(Path(temp_html).resolve().as_uri())
                # 等待 SVG 渲染出来，或者等待错误信息出现
                try:
                    page.wait_for_selector('.mermaid svg', state='attached', timeout=5000)
                    element = page.locator('.mermaid svg')
                    status = "✅ 已渲染 Mermaid"
                except Exception:
                    page.wait_for_selector('#error-container div', state='attached', timeout=5000)
                    element = page.locator('#error-container')
                    status = "⚠️ Mermaid 语法错误，已生成错误提示图"

                page.wait_for_timeout(500) # 等待渲染稳定
                element.screenshot(path=img_path)
                page.close()
                print(f"  {status}: {img_name} (5x超清PNG)")
            except Exception as e:
                print(f"  ❌ 渲染 Mermaid 失败: {e}")
            finally:
                if os.path.exists(temp_html):
                    os.remove(temp_html)
                    
            abs_img_path = os.path.abspath(img_path)
            # 添加 width=100% 确保 Word 中的矢量图也能自适应页面宽度
            return f"\n![架构图表]({abs_img_path}){{width=100%}}\n"
            
        new_content = pattern.sub(repl, content)
        browser.close()
        
    return new_content

def _demote_headings(text, levels=2):
    """把 markdown 中所有标题降 `levels` 级（# -> ###, ## -> ####, …）。
    超过 ###### 的部分截断到 6 级（markdown 上限）。"""
    def repl(m):
        hashes = m.group(1)
        new_hashes = '#' * min(6, len(hashes) + levels)
        return f"{new_hashes} "
    return re.sub(r'^(#{1,6}) ', repl, text, flags=re.MULTILINE)


def _extract_story_title(story_text):
    """从 stories/xx-*.md 中提取首个 `# ` 标题里 `：` 后的文本作为故事标题。
    例如 "# 用户故事：车辆管理员管理车辆档案" -> "车辆管理员管理车辆档案"。
    取不到时返回 None。"""
    m = re.search(r'^# +(.+?)$', story_text, flags=re.MULTILINE)
    if not m:
        return None
    title = m.group(1).strip()
    if '：' in title:
        title = title.split('：', 1)[1].strip()
    elif ':' in title:
        title = title.split(':', 1)[1].strip()
    return title or None


def _extract_model_title(model_text):
    """从 models/xx-*.md 中提取首个 `# ` 标题里 `—`/`-` 前的实体名。
    例如 "# 车辆 — 实体说明" -> "车辆"。"""
    m = re.search(r'^# +(.+?)$', model_text, flags=re.MULTILINE)
    if not m:
        return None
    title = m.group(1).strip()
    for sep in ('—', '–', '-'):
        if sep in title:
            title = title.split(sep, 1)[0].strip()
            break
    return title or None


def _strip_first_h1(text):
    """删除第一行的 `# ` 顶级标题（连带其下方紧邻的空行/引用块）。"""
    lines = text.splitlines()
    out = []
    dropped_h1 = False
    for ln in lines:
        if not dropped_h1 and ln.startswith('# '):
            dropped_h1 = True
            continue
        out.append(ln)
    return '\n'.join(out)


def _strip_manual_toc(text):
    """删除主 PRD 里手写的 `## 目录 ... ---` 整段（含其后的水平分隔线）。
    模板自带真目录，避免在 Word 里出现两份目录。"""
    return re.sub(
        r'## 目录[\s\S]*?(?=^## |\Z)',
        '',
        text,
        count=1,
        flags=re.MULTILINE,
    )


def build_compiled_content(main_content, story_blocks, model_blocks):
    """新拼接策略：

    - 丢弃主 PRD 顶级 H1（与模板封面重复）
    - 丢弃主 PRD 手写目录段（模板自带真目录）
    - 第 4 章"用户故事"就地展开 stories 全文（每个故事 ### 4.x）
    - 第 6 章"核心实体索引"就地展开 models 全文（每个实体 ### 6.x）并改名为
      "核心实体详细说明"
    - 其余章节保留原文（含第 7 章页面清单与附录 A/B/C）

    story_blocks/model_blocks 形如 [(序号, 标题, 原文), ...]，
    其中原文里已经处理过截图替换（stories）/Mermaid 渲染前的中间文本。
    """
    main_content = _strip_first_h1(main_content)
    main_content = _strip_manual_toc(main_content)

    # 4. 用户故事 -> 展开
    if story_blocks:
        body_parts = []
        for idx, title, raw in story_blocks:
            head_title = title or f"故事 {idx}"
            body_parts.append(f"### 4.{idx} 故事{idx}：{head_title}\n")
            body_parts.append(_demote_headings(_strip_first_h1(raw), levels=2))
            body_parts.append("\n")
        new_section = "## 4. 用户故事\n\n" + "\n".join(body_parts).strip() + "\n\n"

        main_content, n = re.subn(
            r'## 4\. 用户故事[\s\S]*?(?=^## \d+\. |^## 附录|\Z)',
            lambda m: new_section,
            main_content,
            count=1,
            flags=re.MULTILINE,
        )
        if n == 0:
            # 主 PRD 没有第 4 章占位，就追加在末尾（极少见兜底）
            main_content += "\n\n" + new_section

    # 6. 核心实体索引 -> 改名为详细说明并展开
    if model_blocks:
        body_parts = []
        for idx, name, raw in model_blocks:
            head_title = name or f"实体 {idx}"
            body_parts.append(f"### 6.{idx} {head_title}\n")
            body_parts.append(_demote_headings(_strip_first_h1(raw), levels=2))
            body_parts.append("\n")
        new_section = (
            "## 6. 核心实体详细说明\n\n" + "\n".join(body_parts).strip() + "\n\n"
        )

        main_content, n = re.subn(
            r'## 6\. 核心实体索引[\s\S]*?(?=^## \d+\. |^## 附录|\Z)',
            lambda m: new_section,
            main_content,
            count=1,
            flags=re.MULTILINE,
        )
        if n == 0:
            main_content += "\n\n" + new_section

    return main_content

def format_beautiful_tables(doc):
    """
    遍历文档中的所有表格，应用统一的美化样式：
    1. 增加 Table Grid 网格边框
    2. 强制表格宽度 100% 撑满页面
    3. 表头第一行增加浅灰背景，文字加粗并居中对齐
    """
    try:
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml, OxmlElement
        
        for table in doc.tables:
            # 基础网格样式
            table.style = 'Table Grid'
            table.autofit = True
            
            # 强制表格宽度为 100%
            tblPr = table._tbl.tblPr
            tblW = OxmlElement('w:tblW')
            tblW.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', '5000')
            tblW.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'pct')
            
            # 移除旧的宽度设置
            for child in tblPr:
                if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW':
                    tblPr.remove(child)
            tblPr.append(tblW)

            # 遍历所有行和单元格，美化表头并修复段落缩进导致的留白问题
            for row_idx, row in enumerate(table.rows):
                is_header = (row_idx == 0)
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    
                    # 单元格垂直居中
                    vAlign = OxmlElement('w:vAlign')
                    vAlign.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'center')
                    tcPr.append(vAlign)

                    if is_header:
                        # 增加浅灰色背景
                        shading_elm = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w')))
                        tcPr.append(shading_elm)

                    # 修复段落样式：清除首行缩进、左缩进和多余的段前段后空白
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.left_indent = 0
                        paragraph.paragraph_format.first_line_indent = 0
                        paragraph.paragraph_format.space_before = 0
                        paragraph.paragraph_format.space_after = 0
                        
                        if is_header:
                            # 表头文字加粗且水平居中
                            for run in paragraph.runs:
                                run.font.bold = True
                            paragraph.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        print(f"⚠️ 表格美化过程发生错误: {e}")

def main(config_path="export-map.yaml"):
    config = load_config(config_path)
    if not config:
        return

    config_dir = os.path.dirname(os.path.abspath(config_path))

    output_dir = os.path.abspath(os.path.join(config_dir, config.get('config', {}).get('output_dir', './output')))
    screenshot_dir = os.path.abspath(os.path.join(config_dir, config.get('config', {}).get('screenshot_dir', './assets/screenshots')))
    output_filename = config.get('config', {}).get('output_filename', '输出文档.docx')
    reference_doc = config.get('config', {}).get('reference_doc')
    project_name = config.get('config', {}).get('project_name')
    prd_dir = config.get('config', {}).get('prd_dir')
    prototype_dir = config.get('config', {}).get('prototype_dir')

    if not prd_dir or not prototype_dir:
        print("错误: 配置文件中必须指定 prd_dir 和 prototype_dir")
        return

    # 转换为绝对路径，完美支持在 YAML 中使用相对路径，并使用 normpath 适配当前操作系统的路径分隔符
    prd_dir = os.path.normpath(os.path.abspath(os.path.join(config_dir, prd_dir)))
    prototype_dir = os.path.normpath(os.path.abspath(os.path.join(config_dir, prototype_dir)))

    # 如果配置文件中没有指定 reference_doc，或者指定的文件不存在，则使用默认模板
    if not reference_doc or not os.path.exists(reference_doc):
        # 获取当前脚本所在目录
        script_dir = os.path.dirname(os.path.abspath(__file__))
        # 寻找默认模板
        default_template = os.path.join(os.path.dirname(script_dir), "templates", "default_template.docx")
        if os.path.exists(default_template):
            reference_doc = default_template
            print(f"提示: 未配置或未找到自定义模板，将使用默认模板: {reference_doc}")
        else:
            print("提示: 未配置自定义模板，且未找到默认模板。生成的文档将没有定制的封面和封底。")
            reference_doc = None

    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(screenshot_dir, exist_ok=True)

    mappings = config.get('mapping', [])
    processed_md_files = []
    md_file_paths = []
    
    for item in mappings:
        md_file = item.get('markdown')
        # 构建 markdown 绝对路径
        abs_md_file = os.path.join(prd_dir, md_file)
        if not os.path.exists(abs_md_file):
            print(f"跳过: 找不到 Markdown 文件 {abs_md_file}")
            continue
            
        if abs_md_file not in md_file_paths:
            md_file_paths.append(abs_md_file)
            
        pages = item.get('pages', [])
        for page in pages:
            html_rel_path = page.get('html')
            if html_rel_path == "手动填写.html":
                continue
                
            html_file = os.path.join(prototype_dir, html_rel_path)
            keyword = page.get('page_name')
            action = page.get('action') # 可能是 None
            
            if not os.path.exists(html_file):
                print(f"跳过: 找不到 HTML 文件 {html_file}")
                continue

            # 生成截图文件名 (如果有动作，文件名加上动作后缀以免覆盖)
            base_name = os.path.splitext(os.path.basename(html_file))[0]
            if action:
                # 清理 action 字符串使其适合做文件名
                safe_action = re.sub(r'[^a-zA-Z0-9]', '_', action)
                screenshot_path = os.path.join(screenshot_dir, f"{base_name}_{safe_action}.png")
            else:
                screenshot_path = os.path.join(screenshot_dir, f"{base_name}.png")

            # 截图
            take_screenshot(html_file, screenshot_path, action)

            # 处理 Markdown (会覆盖更新 .tmp 文件)
            # 对于新结构，我们需要匹配 "### 1.1 页面名称：{keyword}" 这样的标题
            # 为了兼容性，在 process_markdown 内部，正则表达式其实是找这个关键字下方的内容
            # 但更精确的锚点是加上前缀
            exact_keyword = f"页面名称[:：]\\s*{keyword}"
            process_markdown(abs_md_file, html_file, screenshot_path, exact_keyword)
        
    # 对原始 md 文件路径进行排序，确保 01, 02, 03 的顺序
    md_file_paths.sort()
    
    # 收集最终要传给 Pandoc 的文件列表
    # 我们不再直接编译 stories，而是编译总 PRD，总 PRD 中应该包含所有内容
    
    # 获取项目根目录，以便后续处理时定位总 PRD 等文件
    if md_file_paths:
         base_dir = os.path.dirname(os.path.abspath(md_file_paths[0]))
         # 尝试向上一级查找项目根目录（通常 stories 的上一级）
         if os.path.basename(base_dir) == 'stories':
              base_dir = os.path.dirname(base_dir)
    else:
         base_dir = os.path.abspath('.')

    # 查找主 PRD 文件 (寻找以 -PRD.md 结尾，或者包含 PRD 字眼的汇总文件)
    main_prd_path = None
    potential_prds = glob.glob(os.path.join(prd_dir, "*-PRD.md")) + glob.glob(os.path.join(prd_dir, "*PRD.md"))
    if potential_prds:
        main_prd_path = potential_prds[0]
         
    if main_prd_path and os.path.exists(main_prd_path):
        # 创建一个汇总的 Markdown 文件
        compiled_md_path = os.path.join(output_dir, "compiled_prd.md")
        
        with open(main_prd_path, 'r', encoding='utf-8') as f:
            main_content = f.read()

        # 收集 models：按文件名前缀序号排序，提取实体名
        model_blocks = []
        model_files = sorted(glob.glob(os.path.join(prd_dir, "models", "*.md")))
        for i, mf in enumerate(model_files, 1):
            with open(mf, 'r', encoding='utf-8') as f:
                raw = f.read()
            name = _extract_model_title(raw) or os.path.basename(mf)
            model_blocks.append((i, name, raw))

        # 收集 stories：按 export-map 里的 markdown 顺序，使用截图替换后的 .tmp
        story_blocks = []
        for i, md_file in enumerate(md_file_paths, 1):
            temp_md = md_file + ".tmp"
            read_path = temp_md if os.path.exists(temp_md) else md_file
            with open(read_path, 'r', encoding='utf-8') as f:
                raw = f.read()
            title = _extract_story_title(raw) or os.path.basename(md_file)
            story_blocks.append((i, title, raw))

        # 组装最终内容（第 4 章展开 stories，第 6 章展开 models）
        full_content = build_compiled_content(main_content, story_blocks, model_blocks)
        
        # 处理 Mermaid 图表
        full_content = render_mermaid_blocks(full_content, screenshot_dir)
        
        # 写入最终编译的文件
        with open(compiled_md_path, 'w', encoding='utf-8') as f:
            f.write(full_content)
            
        processed_md_files = [compiled_md_path]
    else:
        # 如果没找到总 PRD，退回到只合并故事
        for md_file in md_file_paths:
            temp_md = md_file + ".tmp"
            if os.path.exists(temp_md):
                 processed_md_files.append(temp_md)
            else:
                 processed_md_files.append(md_file)

    if not processed_md_files:
        print("没有处理任何 Markdown 文件，跳过生成 Word。")
        return

    # 使用 pandoc 将 Markdown 生成 Word
    output_docx = os.path.join(output_dir, output_filename)
    temp_pandoc_output = os.path.join(output_dir, "temp_pandoc_output.docx")
    print(f"正在使用 Pandoc 生成 Word 正文: {temp_pandoc_output}")
    
    import shutil
    pandoc_exec = shutil.which("pandoc")
    if not pandoc_exec:
        # 尝试常见路径
        for p in ["/usr/local/bin/pandoc", "/opt/homebrew/bin/pandoc", os.path.expanduser("~/Applications/pandoc/pandoc")]:
            if os.path.exists(p):
                pandoc_exec = p
                break
    if not pandoc_exec:
        print("❌ 未找到 Pandoc，请确保已安装 Pandoc 并将其添加到 PATH 中。")
        return

    pandoc_cmd = [pandoc_exec, processed_md_files[0], "-o", temp_pandoc_output]
    if reference_doc and os.path.exists(reference_doc):
        # 让 Pandoc 使用模板的样式，这样生成的正文就能继承 Heading 1/2 等样式
        pandoc_cmd.extend(["--reference-doc", reference_doc])
        
    try:
        import subprocess
        result = subprocess.run(pandoc_cmd, capture_output=True, text=True)
        if result.returncode != 0:
            print(f"❌ Pandoc 转换失败:\n{result.stderr}")
            return
        print("✅ Pandoc 正文生成成功！")
    except Exception as e:
        print(f"❌ Pandoc 转换发生异常: {e}")
        return

    # 模板注入：使用 docxcompose 将生成的正文插入到参考模板的 {{ content }} 占位符中
    if reference_doc and os.path.exists(reference_doc):
        print(f"尝试将内容无缝注入到模板 {reference_doc} 的 {{{{ content }}}} 占位符中...")
        try:
            from docx import Document
            from docxcompose.composer import Composer
            
            master = Document(reference_doc)

            # 替换模板变量（封面 {{ project_name }} 在文本框内，需走全树遍历）
            if project_name:
                n = _replace_placeholder_everywhere(master.element, '{{ project_name }}', project_name)
                print(f"已替换模板中 {n} 处 {{{{ project_name }}}} 为: {project_name}")

            # 定位 {{ content }} 占位段落在 body 中的索引：正文要插到这里，附录留在后面
            body = master.element.body
            content_index = None
            content_p = None
            for i, child in enumerate(list(body)):
                if child.tag != _W_P:
                    continue
                text = "".join(t.text or "" for t in child.iter(_W_T))
                if '{{ content }}' in text or '{{content}}' in text.replace(" ", ""):
                    content_index = i
                    content_p = child
                    break

            composer = Composer(master)
            content_doc = Document(temp_pandoc_output)
            
            # **执行表格深度美化**：解决 Pandoc 默认 Normal Table 没有边框且排版普通的问题
            print("正在对正文中的表格执行深度美化 (100%宽、边框、表头背景色)...")
            format_beautiful_tables(content_doc)

            # 清掉内容文档里 r:id 缺失的孤立 VML 形状，避免 docxcompose.add_shapes KeyError
            orphans = _strip_orphan_vml_shapes(content_doc)
            if orphans:
                print(f"已清理 {orphans} 个无效 VML 形状（防止 docxcompose 关系迁移报错）")

            # **解决内容重复的核心逻辑**：
            # Pandoc 使用 --reference-doc 生成的 content_doc 会带有模板的封面、目录和最后一页（签字页）。
            # 如果直接把 content_doc 插入到 master 里，就会导致两份封面、两份目录！
            # 因此，在插入前，我们需要把 content_doc 里的非正文部分剔除。
            # Pandoc 插入的正文通常是在原模板的 body 末尾或者覆盖了某些部分，
            # 由于我们很难精确判断哪里是正文，比较稳妥的做法是：
            # 我们不使用 --reference-doc 带来的副作用，或者我们只从 content_doc 中提取内容。
            # 更好的做法是：提供给 Pandoc 一个"干净的" reference-doc（仅含样式，不含封面）。
            pass


            if content_index is not None:
                composer.insert(content_index, content_doc)
                # 插入后占位段落被推后，直接按节点引用删除即可
                content_p.getparent().remove(content_p)
                print(f"✅ 正文已插入到 {{{{ content }}}} 占位符位置（body index={content_index}）")
            else:
                print("⚠️ 模板中未找到 {{ content }} 占位符，正文将追加到模板末尾")
                composer.append(content_doc)

            # 模板里的目录是真正的 TOC 字段；插入正文后字段值仍是旧的，
            # 标记 updateFields=true 让 Word 打开时自动按实际章节刷新目录
            _force_toc_refresh_on_open(master)

            composer.save(output_docx)
            print("🎉 封面、封底和正文融合成功！最终文档已就绪（Word 首次打开会提示更新目录）。")
            
        except ImportError:
            print("⚠️ 缺少 python-docx 或 docxcompose 库，请执行: pip3 install python-docx docxcompose")
        except Exception as e:
            import traceback
            traceback.print_exc()
            print(f"⚠️ 模板融合过程发生错误: {e}")
            print("💡 回退方案: 将保留纯正文格式的文档。")
            # 即使失败，我们将之前生成好的 content_docx (现在变量名叫 output_docx) 作为最终产物
            pass

    # 清理临时文件
    for md_file in md_file_paths:
        temp_md = md_file + ".tmp"
        if os.path.exists(temp_md):
            os.remove(temp_md)
            
    if os.path.exists(compiled_md_path):
        # 可选：是否保留编译后的总 markdown 文件供检查
        # os.remove(compiled_md_path)
        pass

if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1:
        config_path = sys.argv[1]
        main(config_path)
    else:
        main()
